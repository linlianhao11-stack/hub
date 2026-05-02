"""Plan 6 v9 Task 2.2 — 5 个写 tool sentinel 归一化。spec §1.3 v3.4。

验证：handler 入口 x = x or None（空字符串 → None，0 warehouse_id → None）。
所有测试 patch 实际 ORM/persistence 层，捕获传入参数断言归一化结果。
"""
from __future__ import annotations

from unittest.mock import AsyncMock, MagicMock, patch

import pytest


# ============================================================
# generate_contract_draft: shipping_address="" → None
# ============================================================

@pytest.mark.asyncio
async def test_generate_contract_draft_sentinel_empty_to_none(monkeypatch):
    """shipping_address/contact/phone="" 必须归一化成 None — spec §1.3 v3.4。"""
    from hub.agent.tools import generate_tools
    from hub.agent.tools.generate_tools import generate_contract_draft

    # 注入 mock 依赖
    mock_erp = MagicMock()
    mock_erp.get_customer = AsyncMock(return_value={"id": 1, "name": "测试客户"})
    mock_erp.get_account_set = AsyncMock(return_value={
        "id": 1, "name": "测试公司", "company_name": "测试公司",
        "bank_name": "测试银行", "bank_account": "001", "tax_id": "",
    })
    mock_erp.get_product = AsyncMock(return_value={"id": 10, "name": "商品A"})
    mock_sender = MagicMock()
    mock_sender.send_file = AsyncMock()
    generate_tools.set_dependencies(sender=mock_sender, erp=mock_erp)

    # 捕获 ContractDraft.create 参数
    captured_create_kwargs: dict = {}

    async def fake_create(**kwargs):
        captured_create_kwargs.update(kwargs)
        d = MagicMock()
        d.id = 99
        d.status = "generated"
        d.save = AsyncMock()
        return d

    mock_binding = MagicMock()
    mock_binding.channel_userid = "ding-u1"

    try:
        with (
            patch("hub.agent.document.contract.ContractTemplate") as mock_template,
            patch("hub.agent.tools.generate_tools.ContractDraft") as mock_draft,
            patch("hub.agent.tools.generate_tools.ChannelUserBinding") as mock_ch_binding,
        ):
            tqs = MagicMock()
            template_obj = MagicMock()
            template_obj.id = 1
            template_obj.is_active = True
            import io
            from docx import Document
            doc = Document()
            doc.add_paragraph("客户：{{customer_name}}")
            buf = io.BytesIO()
            doc.save(buf)
            import base64
            template_obj.file_storage_key = base64.b64encode(buf.getvalue()).decode()
            template_obj.placeholders = []
            tqs.first = AsyncMock(return_value=template_obj)
            mock_template.filter.return_value = tqs

            mock_draft.create = AsyncMock(side_effect=fake_create)
            mock_draft.filter.return_value.first = AsyncMock(return_value=None)

            bqs = MagicMock()
            bqs.first = AsyncMock(return_value=mock_binding)
            mock_ch_binding.filter.return_value = bqs

            await generate_contract_draft(
                template_id=1,
                customer_id=1,
                items=[{"product_id": 10, "qty": 1, "price": 100.0}],
                shipping_address="",       # 空字符串 — 应归一化成 None
                shipping_contact="",       # 同上
                shipping_phone="",         # 同上
                contract_no="",            # 同上
                payment_terms="",          # 同上
                tax_rate="",               # 同上
                extras={},
                hub_user_id=1,
                conversation_id="conv-sentinel-test",
                acting_as_user_id=1,
            )

        # 断言：空字符串 optional 字段不会被注入 merged_extras 中
        # （handler 内 if shipping_address: 只在非 None/非空时注入 top_level_extras）
        merged = captured_create_kwargs.get("extras", {})
        assert "shipping_address" not in merged, (
            "shipping_address='' 归一化成 None 后不应写入 extras"
        )
        assert "shipping_contact" not in merged, (
            "shipping_contact='' 归一化成 None 后不应写入 extras"
        )
        assert "shipping_phone" not in merged, (
            "shipping_phone='' 归一化成 None 后不应写入 extras"
        )
        assert "contract_no" not in merged, (
            "contract_no='' 归一化成 None 后不应写入 extras"
        )
        assert "payment_terms" not in merged, (
            "payment_terms='' 归一化成 None 后不应写入 extras"
        )
        assert "tax_rate" not in merged, (
            "tax_rate='' 归一化成 None 后不应写入 extras"
        )
    finally:
        generate_tools.set_dependencies(sender=None, erp=None)


# ============================================================
# generate_price_quote: extras={} sentinel（extras or {} 已是 {}）
# ============================================================

@pytest.mark.asyncio
async def test_generate_price_quote_sentinel_empty_extras(monkeypatch):
    """generate_price_quote: extras=None → {} 归一化（通过 generate_contract_draft 验证）。"""
    from hub.agent.tools import generate_tools
    from hub.agent.tools.generate_tools import generate_price_quote

    mock_erp = MagicMock()
    mock_erp.get_customer = AsyncMock(return_value={"id": 1, "name": "测试客户"})
    mock_erp.get_account_set = AsyncMock(return_value={
        "id": 1, "name": "测试公司", "company_name": "测试公司",
        "bank_name": "测试银行", "bank_account": "001", "tax_id": "",
    })
    mock_erp.get_product = AsyncMock(return_value={"id": 10, "name": "商品A"})
    mock_sender = MagicMock()
    mock_sender.send_file = AsyncMock()
    generate_tools.set_dependencies(sender=mock_sender, erp=mock_erp)

    captured: dict = {}

    async def fake_create(**kwargs):
        captured.update(kwargs)
        d = MagicMock()
        d.id = 77
        d.status = "generated"
        d.save = AsyncMock()
        return d

    mock_binding = MagicMock()
    mock_binding.channel_userid = "ding-u2"

    try:
        with (
            patch("hub.agent.tools.generate_tools.ContractTemplate") as mock_tpl_outer,
            patch("hub.agent.document.contract.ContractTemplate") as mock_template,
            patch("hub.agent.tools.generate_tools.ContractDraft") as mock_draft,
            patch("hub.agent.tools.generate_tools.ChannelUserBinding") as mock_ch_binding,
        ):
            # generate_price_quote 先 filter quote 模板
            tpl = MagicMock()
            tpl.id = 2
            outer_qs = MagicMock()
            outer_qs.first = AsyncMock(return_value=tpl)
            mock_tpl_outer.filter.return_value = outer_qs

            import io
            from docx import Document
            doc = Document()
            doc.add_paragraph("报价：{{customer_name}}")
            buf = io.BytesIO()
            doc.save(buf)
            import base64
            inner_tpl = MagicMock()
            inner_tpl.id = 2
            inner_tpl.is_active = True
            inner_tpl.file_storage_key = base64.b64encode(buf.getvalue()).decode()
            inner_tpl.placeholders = []
            inner_qs = MagicMock()
            inner_qs.first = AsyncMock(return_value=inner_tpl)
            mock_template.filter.return_value = inner_qs

            mock_draft.create = AsyncMock(side_effect=fake_create)
            mock_draft.filter.return_value.first = AsyncMock(return_value=None)

            bqs = MagicMock()
            bqs.first = AsyncMock(return_value=mock_binding)
            mock_ch_binding.filter.return_value = bqs

            result = await generate_price_quote(
                customer_id=1,
                items=[{"product_id": 10, "qty": 2, "price": 50.0}],
                extras=None,   # None → 应归一化成 {}
                hub_user_id=1,
                conversation_id="conv-quote-sentinel",
                acting_as_user_id=1,
            )

        # 成功执行（extras=None 不应导致崩溃）
        assert result.get("file_sent") is True
        # extras 应被传成 {} 不是 None（generate_contract_draft 的 safe_extras = extras if dict else {}）
        stored_extras = captured.get("extras", {})
        assert isinstance(stored_extras, dict), "extras 必须是 dict，不能是 None"
    finally:
        generate_tools.set_dependencies(sender=None, erp=None)


# ============================================================
# create_voucher_draft: rule_matched="" → None
# ============================================================

@pytest.mark.asyncio
async def test_create_voucher_draft_sentinel_empty_to_none():
    """rule_matched='' 必须归一化成 None，持久化时 rule_matched=None。"""
    from hub.agent.tools.draft_tools import create_voucher_draft

    captured: dict = {}

    async def fake_create(**kwargs):
        captured.update(kwargs)
        d = MagicMock()
        d.id = 55
        d.status = "pending"
        return d

    voucher_data = {
        "entries": [{"account": "应付账款", "debit": 500, "credit": 0}],
        "total_amount": 500.0,
        "summary": "测试凭证 sentinel",
    }

    with patch("hub.agent.tools.draft_tools.VoucherDraft") as mock_model:
        mock_model.filter.return_value.first = AsyncMock(return_value=None)
        mock_model.create = AsyncMock(side_effect=fake_create)
        # _get_max_voucher_amount 依赖 SystemConfig
        with patch("hub.agent.tools.draft_tools._get_max_voucher_amount", new=AsyncMock(return_value=1_000_000)):
            result = await create_voucher_draft(
                voucher_data=voucher_data,
                rule_matched="",   # 空字符串 — 应归一化成 None
                hub_user_id=1,
                conversation_id="conv-voucher-sentinel",
                acting_as_user_id=1,
                confirmation_action_id="action-sentinel-v",
            )

    assert result["idempotent_replay"] is False
    assert captured["rule_matched"] is None, (
        f"rule_matched='' 应归一化成 None，实际: {captured['rule_matched']!r}"
    )


# ============================================================
# create_price_adjustment_request: reason="" → None
# ============================================================

@pytest.mark.asyncio
async def test_create_price_adjustment_request_sentinel_empty_to_none():
    """reason='' 必须归一化成 None，持久化时 reason=None。"""
    from hub.agent.tools.draft_tools import create_price_adjustment_request, set_erp_adapter

    captured: dict = {}

    async def fake_create(**kwargs):
        captured.update(kwargs)
        r = MagicMock()
        r.id = 66
        r.status = "pending"
        return r

    mock_erp = AsyncMock()
    mock_erp.get_product_customer_prices = AsyncMock(return_value={"items": [{"price": 80.0}]})
    set_erp_adapter(mock_erp)

    try:
        with patch("hub.agent.tools.draft_tools.PriceAdjustmentRequest") as mock_model:
            mock_model.filter.return_value.first = AsyncMock(return_value=None)
            mock_model.create = AsyncMock(side_effect=fake_create)

            result = await create_price_adjustment_request(
                customer_id=10,
                product_id=20,
                new_price=75.0,
                reason="",   # 空字符串 — 应归一化成 None
                hub_user_id=1,
                conversation_id="conv-price-sentinel",
                acting_as_user_id=1,
                confirmation_action_id="action-sentinel-p",
            )

        assert result["idempotent_replay"] is False
        assert captured["reason"] is None, (
            f"reason='' 应归一化成 None，实际: {captured['reason']!r}"
        )
    finally:
        set_erp_adapter(None)


# ============================================================
# create_stock_adjustment_request: reason="" → None, warehouse_id=0 → None
# ============================================================

@pytest.mark.asyncio
async def test_create_stock_adjustment_request_sentinel_empty_to_none():
    """reason='' → None，warehouse_id=0 → None（spec §1.3 v3.4 sentinel 约定）。"""
    from hub.agent.tools.draft_tools import create_stock_adjustment_request, set_erp_adapter

    captured: dict = {}

    async def fake_create(**kwargs):
        captured.update(kwargs)
        r = MagicMock()
        r.id = 77
        r.status = "pending"
        return r

    mock_erp = AsyncMock()
    set_erp_adapter(mock_erp)

    try:
        with patch("hub.agent.tools.draft_tools.StockAdjustmentRequest") as mock_model:
            mock_model.filter.return_value.first = AsyncMock(return_value=None)
            mock_model.create = AsyncMock(side_effect=fake_create)

            result = await create_stock_adjustment_request(
                product_id=30,
                adjustment_qty=5.0,
                reason="",          # 空字符串 — 应归一化成 None
                warehouse_id=0,     # 0 表示"未指定" — 应归一化成 None
                hub_user_id=1,
                conversation_id="conv-stock-sentinel",
                acting_as_user_id=1,
                confirmation_action_id="action-sentinel-s",
            )

        assert result["idempotent_replay"] is False
        assert captured["reason"] is None, (
            f"reason='' 应归一化成 None，实际: {captured['reason']!r}"
        )
        assert captured["warehouse_id"] is None, (
            f"warehouse_id=0 应归一化成 None，实际: {captured['warehouse_id']!r}"
        )
    finally:
        set_erp_adapter(None)
