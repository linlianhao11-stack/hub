"""HUB 后台合同模板管理路由测试（Plan 6 Task 11）。

覆盖场景：
  1. 上传合法 docx + 解析占位符 → 200 + 持久化 + 返 placeholders
  2. 上传非 .docx 文件 → 400
  3. 超过 5MB 文件 → 400
  4. 空文件 → 400
  5. 非法 template_type → 400
  6. 列表返回全部模板
  7. 按 is_active 筛选
  8. 获取占位符（GET /placeholders）
  9. 更新元信息（PUT）
 10. 禁用再启用（disable / enable）
 11. 操作不存在模板 → 404

v2 加固（review C1/C2/I1/M1）新增场景：
 12. 未认证请求 GET list → 401/403（C1）
 13. 未认证请求 GET placeholders → 401/403（C1）
 14. upload 后 AuditLog 持久化（C2）
 15. disable 后 AuditLog 持久化（C2）
 16. 上传 6MB 文件 stream 期间 abort → 400（I1）
 17. 嵌套表格 + 中文占位符都被识别（M1）
"""
from __future__ import annotations

import io
from unittest.mock import AsyncMock

import pytest
import pytest_asyncio
from httpx import ASGITransport, AsyncClient


# ──────────────────────────────────────────────
# 辅助：生成内含占位符的 docx 字节流
# ──────────────────────────────────────────────

def _make_docx_bytes(text: str = "客户：{{customer_name}}\n金额：{{total}}") -> bytes:
    """创建一个最小 docx，包含指定文本（含 {{xxx}} 占位符）。"""
    from docx import Document
    doc = Document()
    for line in text.split("\n"):
        doc.add_paragraph(line)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


_DOCX_CONTENT_TYPE = (
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
)


# ──────────────────────────────────────────────
# 认证 fixture（复用既有 admin 建立风格）
# ──────────────────────────────────────────────

async def _setup_admin(erp_user_id: int = 1, role_code: str = "platform_admin"):
    """构造 platform_admin 身份 + 注入 session_auth，返回 (transport, cookie, hub_user)。"""
    from hub.auth.erp_session import ErpSessionAuth
    from hub.models import DownstreamIdentity, HubRole, HubUser, HubUserRole
    from hub.seed import run_seed
    from main import app

    await run_seed()

    user = await HubUser.create(display_name=f"admin_{erp_user_id}")
    await DownstreamIdentity.create(
        hub_user=user, downstream_type="erp", downstream_user_id=erp_user_id,
    )
    role = await HubRole.get(code=role_code)
    await HubUserRole.create(hub_user_id=user.id, role_id=role.id)

    erp = AsyncMock()
    erp.get_me = AsyncMock(return_value={
        "id": erp_user_id, "username": f"u{erp_user_id}", "permissions": [],
    })
    auth = ErpSessionAuth(erp_adapter=erp)
    app.state.session_auth = auth
    cookie = auth._encode_cookie({
        "jwt": "tok", "user": {"id": erp_user_id, "username": f"u{erp_user_id}"},
    })
    transport = ASGITransport(app=app)
    return transport, cookie, user


@pytest_asyncio.fixture
async def admin_client():
    """已登录的 platform_admin 客户端（含 hub_user）。"""
    transport, cookie, user = await _setup_admin()
    async with AsyncClient(
        transport=transport, base_url="http://t",
        cookies={"hub_session": cookie},
    ) as ac:
        yield ac, user


# ──────────────────────────────────────────────
# 测试用例
# ──────────────────────────────────────────────

BASE = "/hub/v1/admin/contract-templates"


@pytest.mark.asyncio
async def test_upload_template_success(admin_client):
    """上传合法 docx，应返回 200、持久化模板、解析出正确占位符数量。"""
    from hub.models.contract import ContractTemplate

    ac, _ = admin_client
    docx_bytes = _make_docx_bytes("客户：{{customer_name}}\n金额：{{total}}")
    resp = await ac.post(
        BASE,
        files={"file": ("合同模板.docx", docx_bytes, _DOCX_CONTENT_TYPE)},
        data={"name": "销售合同模板", "template_type": "sales", "description": "测试用"},
    )
    assert resp.status_code == 200, resp.text
    body = resp.json()
    assert body["id"] is not None
    assert body["name"] == "销售合同模板"
    assert len(body["placeholders"]) == 2
    names = {p["name"] for p in body["placeholders"]}
    assert names == {"customer_name", "total"}

    # 持久化验证
    record = await ContractTemplate.filter(id=body["id"]).first()
    assert record is not None
    assert record.name == "销售合同模板"
    assert record.is_active is True


@pytest.mark.asyncio
async def test_upload_invalid_extension_rejected(admin_client):
    """传非 .docx 文件 → 400。"""
    ac, _ = admin_client
    resp = await ac.post(
        BASE,
        files={"file": ("模板.pdf", b"fake pdf content", "application/pdf")},
        data={"name": "pdf测试", "template_type": "sales"},
    )
    assert resp.status_code == 400
    assert "docx" in resp.json()["detail"]


@pytest.mark.asyncio
async def test_upload_oversized_rejected(admin_client):
    """上传超过 5MB 的文件 → 400，错误信息提及 5MB。"""
    ac, _ = admin_client
    # 先构造一个合法 docx 外壳（小），然后在 data 里模拟超大字节（直接传 6MB 的字节流）
    big_bytes = b"0" * (5 * 1024 * 1024 + 1)
    resp = await ac.post(
        BASE,
        files={"file": ("大文件.docx", big_bytes, _DOCX_CONTENT_TYPE)},
        data={"name": "大文件测试", "template_type": "sales"},
    )
    assert resp.status_code == 400
    assert "5MB" in resp.json()["detail"]


@pytest.mark.asyncio
async def test_upload_empty_file_rejected(admin_client):
    """上传空 bytes → 400。"""
    ac, _ = admin_client
    resp = await ac.post(
        BASE,
        files={"file": ("空文件.docx", b"", _DOCX_CONTENT_TYPE)},
        data={"name": "空文件", "template_type": "sales"},
    )
    assert resp.status_code == 400


@pytest.mark.asyncio
async def test_upload_invalid_template_type(admin_client):
    """传 template_type='xxx' → 400。"""
    ac, _ = admin_client
    docx_bytes = _make_docx_bytes("内容")
    resp = await ac.post(
        BASE,
        files={"file": ("模板.docx", docx_bytes, _DOCX_CONTENT_TYPE)},
        data={"name": "类型错误测试", "template_type": "invalid_type"},
    )
    assert resp.status_code == 400


@pytest.mark.asyncio
async def test_list_templates_returns_all(admin_client):
    """插入 3 条，GET 列表返回全部 3 条。"""
    from hub.models.contract import ContractTemplate

    ac, user = admin_client
    for i in range(3):
        await ContractTemplate.create(
            name=f"模板{i}",
            template_type="sales",
            file_storage_key="dummy",
            placeholders=[],
            is_active=True,
            created_by_hub_user_id=user.id,
        )

    resp = await ac.get(BASE)
    assert resp.status_code == 200
    body = resp.json()
    assert body["total"] >= 3
    assert len(body["items"]) >= 3


@pytest.mark.asyncio
async def test_list_filter_by_active(admin_client):
    """is_active=true 只返回启用状态的模板。"""
    from hub.models.contract import ContractTemplate

    ac, user = admin_client
    await ContractTemplate.create(
        name="启用的模板", template_type="sales", file_storage_key="k1",
        placeholders=[], is_active=True, created_by_hub_user_id=user.id,
    )
    await ContractTemplate.create(
        name="禁用的模板", template_type="sales", file_storage_key="k2",
        placeholders=[], is_active=False, created_by_hub_user_id=user.id,
    )

    resp = await ac.get(BASE, params={"is_active": "true"})
    assert resp.status_code == 200
    body = resp.json()
    assert all(item["is_active"] for item in body["items"])

    resp2 = await ac.get(BASE, params={"is_active": "false"})
    assert resp2.status_code == 200
    body2 = resp2.json()
    assert all(not item["is_active"] for item in body2["items"])


@pytest.mark.asyncio
async def test_get_placeholders(admin_client):
    """GET /{id}/placeholders 返回该模板的占位符列表。"""
    from hub.models.contract import ContractTemplate

    ac, user = admin_client
    placeholders = [
        {"name": "buyer", "type": "string", "required": True},
        {"name": "amount", "type": "string", "required": True},
    ]
    tpl = await ContractTemplate.create(
        name="占位符测试模板", template_type="purchase", file_storage_key="k",
        placeholders=placeholders, is_active=True, created_by_hub_user_id=user.id,
    )

    resp = await ac.get(f"{BASE}/{tpl.id}/placeholders")
    assert resp.status_code == 200
    body = resp.json()
    assert len(body["placeholders"]) == 2
    names = {p["name"] for p in body["placeholders"]}
    assert names == {"buyer", "amount"}


@pytest.mark.asyncio
async def test_update_template_metadata(admin_client):
    """PUT 更新 name + description，持久化正确。"""
    from hub.models.contract import ContractTemplate

    ac, user = admin_client
    tpl = await ContractTemplate.create(
        name="原始名称", template_type="sales", file_storage_key="k",
        placeholders=[], is_active=True, created_by_hub_user_id=user.id,
    )

    resp = await ac.put(
        f"{BASE}/{tpl.id}",
        json={"name": "更新后名称", "description": "新描述"},
    )
    assert resp.status_code == 200

    await tpl.refresh_from_db()
    assert tpl.name == "更新后名称"
    assert tpl.description == "新描述"


@pytest.mark.asyncio
async def test_disable_then_enable(admin_client):
    """disable → is_active=False；enable → is_active=True；幂等操作不报错。"""
    from hub.models.contract import ContractTemplate

    ac, user = admin_client
    tpl = await ContractTemplate.create(
        name="启用禁用测试", template_type="framework", file_storage_key="k",
        placeholders=[], is_active=True, created_by_hub_user_id=user.id,
    )

    # 禁用
    resp = await ac.post(f"{BASE}/{tpl.id}/disable")
    assert resp.status_code == 200
    await tpl.refresh_from_db()
    assert tpl.is_active is False

    # 再次禁用（幂等）
    resp2 = await ac.post(f"{BASE}/{tpl.id}/disable")
    assert resp2.status_code == 200

    # 启用
    resp3 = await ac.post(f"{BASE}/{tpl.id}/enable")
    assert resp3.status_code == 200
    await tpl.refresh_from_db()
    assert tpl.is_active is True

    # 再次启用（幂等）
    resp4 = await ac.post(f"{BASE}/{tpl.id}/enable")
    assert resp4.status_code == 200


@pytest.mark.asyncio
async def test_disable_nonexistent_404(admin_client):
    """对不存在的模板 ID 调用 disable → 404。"""
    ac, _ = admin_client
    resp = await ac.post(f"{BASE}/999999/disable")
    assert resp.status_code == 404


# ──────────────────────────────────────────────
# v2 加固测试（review C1/C2/I1/M1）
# ──────────────────────────────────────────────

@pytest.mark.asyncio
async def test_list_requires_read_permission():
    """v2 加固（review C1）：未登录用户不能列出模板，应返回 401/403。"""
    from main import app
    transport = ASGITransport(app=app)
    async with AsyncClient(transport=transport, base_url="http://t") as ac:
        resp = await ac.get(BASE)
    assert resp.status_code in (401, 403)


@pytest.mark.asyncio
async def test_get_placeholders_requires_read_permission():
    """v2 加固（review C1）：未登录用户不能查看占位符，应返回 401/403。"""
    from main import app
    transport = ASGITransport(app=app)
    async with AsyncClient(transport=transport, base_url="http://t") as ac:
        resp = await ac.get(f"{BASE}/1/placeholders")
    assert resp.status_code in (401, 403)


@pytest.mark.asyncio
async def test_upload_writes_audit_log(admin_client):
    """v2 加固（review C2）：upload 后 AuditLog 持久化，含正确字段。"""
    from hub.models.audit import AuditLog

    ac, _ = admin_client
    docx_bytes = _make_docx_bytes("客户：{{customer_name}}\n金额：{{total}}")
    resp = await ac.post(
        BASE,
        files={"file": ("审计测试.docx", docx_bytes, _DOCX_CONTENT_TYPE)},
        data={"name": "审计测试模板", "template_type": "sales", "description": ""},
    )
    assert resp.status_code == 200, resp.text
    created_id = resp.json()["id"]

    log = await AuditLog.filter(
        action="upload_contract_template",
        target_id=str(created_id),
    ).first()
    assert log is not None
    assert log.target_type == "contract_template"
    assert log.detail.get("placeholders_count") == 2
    assert log.detail.get("name") == "审计测试模板"


@pytest.mark.asyncio
async def test_disable_writes_audit_log(admin_client):
    """v2 加固（review C2）：disable 后 AuditLog 持久化，含正确字段。"""
    from hub.models.audit import AuditLog
    from hub.models.contract import ContractTemplate

    ac, user = admin_client
    tpl = await ContractTemplate.create(
        name="禁用审计测试", template_type="sales", file_storage_key="k",
        placeholders=[], is_active=True, created_by_hub_user_id=user.id,
    )

    resp = await ac.post(f"{BASE}/{tpl.id}/disable")
    assert resp.status_code == 200

    log = await AuditLog.filter(
        action="disable_contract_template",
        target_id=str(tpl.id),
    ).first()
    assert log is not None
    assert log.target_type == "contract_template"
    assert log.detail.get("name") == "禁用审计测试"


@pytest.mark.asyncio
async def test_upload_oversized_streams_aborted(admin_client):
    """v2 加固（review I1）：上传 6MB 文件应在 stream 期间提前中断，返回 400。"""
    ac, _ = admin_client
    big_bytes = b"\x00" * (6 * 1024 * 1024)
    resp = await ac.post(
        BASE,
        files={"file": ("big.docx", big_bytes, _DOCX_CONTENT_TYPE)},
        data={"name": "超大文件", "template_type": "sales"},
    )
    assert resp.status_code == 400
    detail = resp.json()["detail"]
    assert "5MB" in detail or "上限" in detail


@pytest.mark.asyncio
async def test_extract_placeholders_nested_tables_and_chinese_names():
    """v2 加固（review M1）：嵌套表格 + 中文占位符名都被正确识别。"""
    from docx import Document

    from hub.routers.admin.contract_templates import _extract_placeholders

    doc = Document()
    doc.add_paragraph("{{客户名}}")
    table = doc.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "{{表格内占位}}"
    # 嵌套表格
    nested = table.rows[1].cells[1].add_table(rows=1, cols=1)
    nested.rows[0].cells[0].text = "{{嵌套占位}}"

    buf = io.BytesIO()
    doc.save(buf)

    result = _extract_placeholders(buf.getvalue())
    names = {p["name"] for p in result}
    assert "客户名" in names
    assert "表格内占位" in names
    assert "嵌套占位" in names
