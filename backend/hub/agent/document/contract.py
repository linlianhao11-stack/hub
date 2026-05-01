"""Plan 6 Task 7：合同模板 docx 渲染。

ContractTemplate.placeholders 形如：
[
  {"name": "customer_name", "type": "string", "required": true},
  {"name": "items_table", "type": "table", "required": true},
  {"name": "total_amount", "type": "decimal", "required": true},
]

模板 docx 中用 {{customer_name}} 标记文本占位；
表格用整行 {{items_table}}（找到所在 cell 后展开 list[dict]）。
"""
from __future__ import annotations

import base64
import copy
import io
import logging
import re
from datetime import UTC, date, datetime

from docx import Document  # python-docx
from docx.oxml.ns import qn
from docx.shared import Cm
from docx.table import _Row

from hub.models.contract import ContractTemplate

logger = logging.getLogger("hub.agent.document.contract")


# v8 staging review #11：模板列宽自动调节（合同 + 报价共用算法，各自规则）
# 凭证 / 调价 / 库存调整等暂未涉及 docx 渲染（落 ERP 表）；
# 后续如有 docx 类型按这个 pattern 加新 columns_config 即可。
SALES_CONTRACT_COLUMNS: list[dict] = [
    {"key": "_idx",     "label": "序号",       "min": 0.7, "max": 1.0, "weight": 1},
    {"key": "name",     "label": "产品名称",   "min": 2.5, "max": 6.0, "weight": 5},
    {"key": "spec",     "label": "规格",       "min": 1.0, "max": 2.0, "weight": 1},
    {"key": "color",    "label": "颜色",       "min": 1.0, "max": 2.0, "weight": 1},
    {"key": "qty",      "label": "数量",       "min": 0.8, "max": 1.5, "weight": 1},
    {"key": "price",    "label": "单价（元）", "min": 1.5, "max": 2.5, "weight": 2},
    {"key": "subtotal", "label": "含税金额",   "min": 1.8, "max": 3.5, "weight": 2},
    {"key": "remark",   "label": "备注",       "min": 1.5, "max": 3.0, "weight": 1},
]

# 报价单（5 列：序号 / 产品名称 / 数量 / 单价 / 小计）
QUOTE_COLUMNS: list[dict] = [
    {"key": "_idx",     "label": "序号",     "min": 0.7, "max": 1.0, "weight": 1},
    {"key": "name",     "label": "产品名称", "min": 4.0, "max": 8.0, "weight": 5},
    {"key": "qty",      "label": "数量",     "min": 1.0, "max": 1.8, "weight": 1},
    {"key": "price",    "label": "单价",     "min": 2.0, "max": 3.5, "weight": 2},
    {"key": "subtotal", "label": "小计",     "min": 2.0, "max": 3.5, "weight": 2},
]

# 模板类型 → columns_config 注册表（统一入口）
_TEMPLATE_COLUMNS_REGISTRY: dict[str, list[dict]] = {
    "sales": SALES_CONTRACT_COLUMNS,
    "quote": QUOTE_COLUMNS,
    # 凭证 / 调价 / 库存调整等如需 docx 化，加这里
    # "voucher_doc": VOUCHER_COLUMNS,
    # "price_adjust_doc": PRICE_ADJUST_COLUMNS,
}


def _estimate_text_width_cm(text, font_size_pt: int = 10) -> float:
    """估算文本在 docx cell 里的实际显示宽度（cm）。

    宋体 10pt 经验值：
    - 中文字符（含全角符号）≈ 0.5 cm
    - ASCII / 数字（半角）≈ 0.2 cm
    """
    s = str(text or "")
    cn_count = sum(
        1 for c in s
        if "一" <= c <= "鿿"        # CJK Unified
        or "　" <= c <= "〿"        # CJK Symbols
        or "＀" <= c <= "￯"        # 全角符号
    )
    other_count = len(s) - cn_count
    # 字号缩放：基于 10pt
    scale = font_size_pt / 10.0
    return (cn_count * 0.5 + other_count * 0.2) * scale


def _calc_column_widths(
    items: list[dict],
    columns_config: list[dict],
    total_cm: float = 14.5,
    padding_cm: float = 0.4,
) -> list[float]:
    """根据 items 实际数据动态算每列宽度（cm）。

    Args:
        items: list of dict，每个 dict 含 columns_config 里 key 对应的字段
        columns_config: [{key, label, min, max, weight}, ...]
        total_cm: 表格总宽（A4 portrait 内容区 ≈ 16cm，留 1.5cm 边给 14.5cm）
        padding_cm: 每列内容边距留白

    Algorithm:
      1. 每列 ideal_width = max(表头宽, 所有 items 该列内容宽) + padding，
         clamp 到 [min_cm, max_cm]
      2. 总宽 ≤ total_cm → 直接用
      3. 总宽 > total_cm → 按 weight 反比缩（weight 大列优先保宽，
         weight=1 列扣得多，weight=5 列扣得少），不低于 min
    """
    n = len(columns_config)
    if n == 0:
        return []

    ideal: list[float] = []
    for col in columns_config:
        max_text = col["label"]  # 表头作为基线
        for it in items or []:
            if not isinstance(it, dict):
                continue
            v = str(it.get(col["key"], "") or "")
            if len(v) > len(max_text):
                max_text = v
        w = _estimate_text_width_cm(max_text) + padding_cm
        ideal.append(max(col["min"], min(col["max"], w)))

    total = sum(ideal)
    if total <= total_cm:
        return ideal

    # 超 total → 按 weight 反比缩（weight 大列保得多）
    excess = total - total_cm
    weights = [col["weight"] for col in columns_config]
    inv_weights = [1.0 / w for w in weights]
    sum_inv = sum(inv_weights) or 1.0

    widths = list(ideal)
    for i, col in enumerate(columns_config):
        share = (inv_weights[i] / sum_inv) * excess
        widths[i] = max(col["min"], widths[i] - share)

    # min 卡住时再扫一遍从 weight 最小列扣
    iter_count = 0
    while sum(widths) > total_cm + 0.01 and iter_count < 8:
        iter_count += 1
        cur_excess = sum(widths) - total_cm
        candidates = [
            (i, col["weight"])
            for i, col in enumerate(columns_config)
            if widths[i] > col["min"] + 0.05
        ]
        if not candidates:
            break  # 都到 min 了，认了
        idx_to_shrink = min(candidates, key=lambda x: x[1])[0]
        col = columns_config[idx_to_shrink]
        room = widths[idx_to_shrink] - col["min"]
        delta = min(cur_excess, room)
        widths[idx_to_shrink] -= delta

    return widths


class TemplateNotFoundError(Exception):
    """合同模板不存在或未启用。"""


class TemplateRenderError(Exception):
    """模板渲染失败（占位符缺失 / 数据格式错 / file_storage_key 格式异常）。"""


# 中文大写人民币转换（v8 staging review #8）
_CN_DIGITS = "零壹贰叁肆伍陆柒捌玖"
_CN_UNITS_4 = ["", "拾", "佰", "仟"]
_CN_GROUPS = ["", "万", "亿", "兆"]


def _yuan_to_chinese(amount) -> str:
    """金额（元，可含 2 位小数 = 角分）转中文大写人民币写法。
    例：
        1234.56 → 壹仟贰佰叁拾肆元伍角陆分
        80000   → 捌万元整
        0       → 零元整
        12.5    → 壹拾贰元伍角
    """
    try:
        amt = float(amount)
    except (TypeError, ValueError):
        return ""

    if amt < 0:
        return "（负）" + _yuan_to_chinese(-amt)

    yuan = int(amt)
    fen_total = round((amt - yuan) * 100)
    if fen_total >= 100:  # round 上溢
        yuan += 1
        fen_total = 0
    jiao = fen_total // 10
    fen = fen_total % 10

    # 整数部分按每 4 位分组
    if yuan == 0:
        yuan_str = "零"
    else:
        groups = []
        n = yuan
        while n > 0:
            groups.append(n % 10000)
            n //= 10000
        parts: list[str] = []
        for gi in range(len(groups) - 1, -1, -1):
            g = groups[gi]
            if g == 0:
                if parts and not parts[-1].endswith("零"):
                    parts.append("零")
                continue
            seg = ""
            zero_pending = False
            for d_idx in range(3, -1, -1):
                digit = (g // (10 ** d_idx)) % 10
                if digit == 0:
                    if seg and not seg.endswith("零"):
                        zero_pending = True
                else:
                    if zero_pending:
                        seg += "零"
                        zero_pending = False
                    seg += _CN_DIGITS[digit] + _CN_UNITS_4[d_idx]
            parts.append(seg + _CN_GROUPS[gi])
        yuan_str = "".join(parts).rstrip("零")

    if jiao == 0 and fen == 0:
        return f"{yuan_str}元整"

    result = f"{yuan_str}元"
    if jiao > 0:
        result += _CN_DIGITS[jiao] + "角"
    elif fen > 0:
        result += "零"  # 元后无角有分要补"零"
    if fen > 0:
        result += _CN_DIGITS[fen] + "分"
    return result


class ContractRenderer:
    """合同模板渲染：从 ContractTemplate 加载 docx → 替换占位符 → 返 bytes。

    模板设计规范（Plan 11 admin 上传模板时遵守）：
    - 占位符 {{xxx}} 必须独立成段落或独立成 cell，不要与其它格式 run 混用
    - 例：✅ 单独段落 "客户：{{customer_name}}"
    - 例：❌ 同一段落 "客户：{{customer_name}}（**重要**）" — bold 格式会丢失
    第一版段落级 + 表格级两级替换；多 run 格式合并问题待 follow-up
    （需"找连续含 {{ 的 run 串、合并文本、保留首 run 格式"逻辑）。
    """

    async def render(
        self,
        *,
        template_id: int,
        customer: dict,
        items: list[dict],
        extras: dict | None = None,
    ) -> bytes:
        """渲染合同 docx。

        Args:
            template_id: ContractTemplate.id
            customer: {"id", "name", "address", ...}（ERP 客户字段）
            items: [{"product_id", "name", "qty", "price", "subtotal"}, ...]
            extras: 额外字段如 contract_no / payment_terms

        Returns:
            docx 字节流（未加密；调用方决定是否存储加密）

        Raises:
            TemplateNotFoundError: template_id 不存在或未启用
            TemplateRenderError: 占位符缺失或数据格式错
        """
        template = await ContractTemplate.filter(
            id=template_id, is_active=True,
        ).first()
        if not template:
            raise TemplateNotFoundError(f"合同模板 {template_id} 不存在或未启用")

        if not template.file_storage_key:
            raise TemplateRenderError(f"模板 {template_id} 缺 file_storage_key")

        template_bytes = await self._load_template_bytes(template)

        try:
            doc = Document(io.BytesIO(template_bytes))
            ctx = self._build_context(template, customer, items, extras or {})

            # v8 staging review #11：根据模板类型 + items 实际内容动态算列宽
            # （只对注册了 columns_config 的类型生效——sales/quote）
            columns_config = _TEMPLATE_COLUMNS_REGISTRY.get(template.template_type)
            column_widths = (
                _calc_column_widths(items or [], columns_config)
                if columns_config else None
            )
            if column_widths:
                logger.info(
                    "合同模板 %s (type=%s) 动态列宽 cm: %s",
                    template.id, template.template_type,
                    [round(w, 2) for w in column_widths],
                )

            self._replace_paragraphs(doc, ctx)
            self._replace_tables(doc, ctx, column_widths=column_widths)
            # 扫描渲染后仍含 {{xxx}} 的占位符（模板 typo 或数据缺字段）
            unknown_placeholders: set[str] = set()
            for para in doc.paragraphs:
                for m in re.finditer(r"\{\{(\w+)\}\}", para.text):
                    unknown_placeholders.add(m.group(1))
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for m in re.finditer(r"\{\{(\w+)\}\}", cell.text):
                            unknown_placeholders.add(m.group(1))
            if unknown_placeholders:
                logger.warning(
                    "合同模板 %s 渲染后仍含未知占位符: %s（可能是模板 typo）",
                    template.id, sorted(unknown_placeholders),
                )
            output = io.BytesIO()
            doc.save(output)
            return output.getvalue()
        except (TemplateRenderError, TemplateNotFoundError):
            raise
        except KeyError as e:
            raise TemplateRenderError(f"模板缺占位符 {e}") from e
        except Exception as e:
            raise TemplateRenderError(f"渲染失败: {e}") from e

    @staticmethod
    async def _load_template_bytes(template: ContractTemplate) -> bytes:
        """加载模板 docx 字节流。

        第一版：template.file_storage_key 存的是 base64 编码的 docx bytes。
        生产可换成 OSS URL / 文件路径。
        """
        try:
            return base64.b64decode(template.file_storage_key)
        except Exception as e:
            raise TemplateRenderError(f"模板 file_storage_key 格式异常: {e}") from e

    @staticmethod
    def _build_context(
        template: ContractTemplate,
        customer: dict,
        items: list[dict],
        extras: dict,
    ) -> dict:
        """组装占位符上下文。

        v2 staging review #5：placeholders 自动识别全部 required=True 太严苛——
        模板里有 18 个占位符（甲方/乙方/订单/收货全套），LLM 不可能每个都传齐。
        修：缺数据时**注入空字符串**而不是抛错，让 docx 渲染产物里
        留白让用户后期手填，至少合同主体（客户/商品/总额）能正确生成。
        """
        # extras 类型加固（防 LLM 传 string 等错型）
        if not isinstance(extras, dict):
            extras = {}

        # v8 review #8：自动算合计金额 + 中文大写（不依赖 LLM 传 total_amount_cn）
        total_amount_num = sum(
            float(i.get("subtotal") or float(i.get("price", 0)) * float(i.get("qty", 0)))
            for i in items
        )

        ctx: dict = {
            # 乙方（来自 ERP customer 详情）
            "customer_name": customer.get("name") or "",
            "customer_address": customer.get("address") or "",
            "customer_id": str(customer.get("id") or ""),
            "customer_phone": customer.get("phone") or "",
            "customer_tax_id": customer.get("tax_id") or "",
            "customer_bank_name": customer.get("bank_name") or "",
            "customer_bank_account": customer.get("bank_account") or "",
            "customer_contact_person": customer.get("contact_person") or "",
            # 订单
            "today": date.today().isoformat(),
            "now": datetime.now(UTC).isoformat(),
            "items": items,
            "total_amount": f"{total_amount_num:,.2f}".rstrip("0").rstrip(".") or "0",
            "total_amount_cn": _yuan_to_chinese(total_amount_num),
            # extras 覆盖前面（甲方账套字段 / shipping / tax_rate 等都从这里来）
            # 但 total_amount / total_amount_cn LLM 传也会被默认覆盖（不允许 LLM 改总额）
            **{k: str(v) if not isinstance(v, (list, dict)) else v for k, v in extras.items()},
        }
        # extras 不允许覆盖系统计算的合计（防 LLM 自己写错金额）
        ctx["total_amount"] = f"{total_amount_num:,.2f}".rstrip("0").rstrip(".") or "0"
        ctx["total_amount_cn"] = _yuan_to_chinese(total_amount_num)

        # required 字段缺数据时注入空字符串而不是抛错
        # 模板里残留的 {{xxx}} 占位符会被替换为空，docx 主体仍可生成（客户/商品/总额都对）
        # 后期 admin / 销售可在 docx 里手填空白字段
        missing_required: list[str] = []
        for ph in (template.placeholders or []):
            if not isinstance(ph, dict):
                continue
            name = ph.get("name")
            if not name:
                continue
            if ph.get("required") and name not in ctx:
                ctx[name] = ""  # 空串占位
                missing_required.append(name)
        if missing_required:
            import logging
            logging.getLogger(__name__).warning(
                "合同模板 %s 渲染时缺数据字段 %s（已用空串兜底，docx 留白）",
                template.id, missing_required,
            )

        return ctx

    @staticmethod
    def _replace_paragraphs(doc: Document, ctx: dict) -> None:
        """段落级 {{key}} 替换。"""
        for para in doc.paragraphs:
            full_text = para.text
            if "{{" not in full_text:
                continue
            new_text = full_text
            for key, value in ctx.items():
                if isinstance(value, (list, dict)):
                    continue  # 列表/字典占位符走 _replace_tables
                placeholder = f"{{{{{key}}}}}"
                if placeholder in new_text:
                    new_text = new_text.replace(placeholder, str(value))
            if new_text != full_text:
                # 清空所有 run，只用第一个 run 写文本
                for run in para.runs:
                    run.text = ""
                if para.runs:
                    para.runs[0].text = new_text
                else:
                    para.add_run(new_text)

    @staticmethod
    def _replace_tables(
        doc: Document, ctx: dict, *, column_widths: list[float] | None = None,
    ) -> None:
        """表格级 {{items_table}} 占位符 → 展开成多行；其余 cell 内占位符替换。

        merged cell 防重复处理：merged 后多 row 共享同一个 XML element（cell._tc），
        用 element id 去重。**不能用 id(cell)** —— Python GC 会复用对象 id，
        导致正常的不同 cell 也被错认为"已处理"（v8 staging review #6 实测踩坑）。
        """
        # v4 加固（v8 staging review #7）：items_table 真正按列展开成 N 行表格，
        # 而不是把所有数据挤到第一个 cell。
        # 模板规则：含 `{{items_table}}` 的整行视为"商品行模板"，期望 8 列：
        #   序号 / 产品名称 / 规格 / 颜色 / 数量 / 单价 / 含税金额 / 备注
        # 渲染流程：在该 row 前插入 N 行（每个 item 一行），最后删除占位 row。
        # 兼容：cell 数 < 8 时按 cell 顺序填前 N 个字段，剩余字段忽略。

        # 第 1 阶段：找到所有"items_table 模板行"，记录 (table, tr_element) 待展开
        items_template_trs: list[tuple] = []
        for table in doc.tables:
            for row in table.rows:
                row_text = " ".join(c.text for c in row.cells)
                if "{{items_table}}" in row_text:
                    items_template_trs.append((table, row, row._tr))
                    break  # 同一 table 只取第一个 items_table 行

        # 第 1.5 阶段（v8 review #11）：动态调整 items_table 所在表格的列宽
        # column_widths 从 render() 算好传进来；不为空时覆盖模板原列宽
        if column_widths:
            for table, _row, _tr in items_template_trs:
                ContractRenderer._set_table_column_widths(table, column_widths)

        # 第 2 阶段：按记录展开
        items = ctx.get("items") or []
        for table, row, tr in items_template_trs:
            ContractRenderer._expand_items_row(table, row, tr, items)

        # 第 3 阶段：扫所有表格 cell 做普通占位符替换（替换幂等，无需去重）
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        para_text = para.text
                        if "{{" not in para_text:
                            continue
                        new_text = para_text
                        for key, value in ctx.items():
                            if isinstance(value, (list, dict)):
                                continue
                            placeholder = f"{{{{{key}}}}}"
                            if placeholder in new_text:
                                new_text = new_text.replace(placeholder, str(value))
                        if new_text != para_text:
                            for run in para.runs:
                                run.text = ""
                            if para.runs:
                                para.runs[0].text = new_text
                            else:
                                para.add_run(new_text)

    @staticmethod
    def _expand_items_row(table, template_row, template_tr, items: list[dict]) -> None:
        """把 items_table 占位行展开成 N 行。

        - 复制模板行的 XML element 作为新行（保留边框 / 字号等格式）
        - 按列 0-7 填：序号 / 产品名称 / 规格 / 颜色 / 数量 / 单价 / 含税金额 / 备注
        - 在原 tr 前插入新行，最后删除原 tr
        """
        parent = template_tr.getparent()
        if parent is None:
            return  # 模板行已脱离 doc，啥都不做

        if not items:
            # 无 items：清空占位行内容（保留行结构）
            for cell in template_row.cells:
                ContractRenderer._set_cell_text(cell, "")
            return

        idx = list(parent).index(template_tr)
        # 8 列字段顺序（与模板对齐；cell 数少于 8 时按顺序填前 N 个）
        for i, item in enumerate(items):
            new_tr = copy.deepcopy(template_tr)
            new_row = _Row(new_tr, table)
            cells = new_row.cells

            qty = item.get("qty") or 0
            price = item.get("price") or 0
            subtotal = item.get("subtotal")
            if subtotal is None:
                try:
                    subtotal = float(price) * float(qty)
                except (TypeError, ValueError):
                    subtotal = ""

            col_values = [
                str(i + 1),                            # 序号
                str(item.get("name") or ""),           # 产品名称
                str(item.get("spec") or ""),           # 规格
                str(item.get("color") or ""),          # 颜色
                str(qty) if qty else "",               # 数量
                str(price) if price else "",           # 单价
                str(subtotal) if subtotal != "" else "",  # 含税金额
                str(item.get("remark") or ""),         # 备注
            ]
            for col_i, value in enumerate(col_values):
                if col_i >= len(cells):
                    break
                ContractRenderer._set_cell_text(cells[col_i], value)

            parent.insert(idx + i, new_tr)

        # 删除原占位行（在所有新行已插入后）
        parent.remove(template_tr)

    @staticmethod
    def _set_cell_text(cell, text: str) -> None:
        """清空 cell 所有 paragraph 的 run，把 text 写到第一个 paragraph 的第一个 run。
        保留 cell 原有格式（边框 / 字号 / 对齐）。"""
        # 清空所有 run text（不删 paragraph，保字号/对齐设置）
        for p in cell.paragraphs:
            for r in p.runs:
                r.text = ""
        if not cell.paragraphs:
            cell.add_paragraph(text)
            return
        p = cell.paragraphs[0]
        if p.runs:
            p.runs[0].text = text
        else:
            p.add_run(text)

    @staticmethod
    def _set_table_column_widths(table, widths_cm: list[float]) -> None:
        """覆盖表格列宽（v8 staging review #11）。

        改 2 处：
          1. <w:tblGrid>/<w:gridCol> ：表格列定义（所有 row 默认遵从）
          2. 每行每 cell 的 width ：确保 Word 严格按新列宽渲染

        merged cell 也兼容：python-docx 的 row.cells 对 merged cell 按 grid index
        返回多次（每个被合并的 grid 位置都返回），按 ci 索引设 width 即可。
        """
        # 1. 修 tblGrid（表格列定义）
        tblGrid = table._tbl.find(qn("w:tblGrid"))
        if tblGrid is not None:
            cols = tblGrid.findall(qn("w:gridCol"))
            for i, col_el in enumerate(cols):
                if i < len(widths_cm):
                    # cm → twips（1 cm = 567 twips）
                    col_el.set(qn("w:w"), str(int(widths_cm[i] * 567)))
        # 2. 每行每 cell 设 width（保险措施，确保 Word 严格遵循）
        for row in table.rows:
            for ci, cell in enumerate(row.cells):
                if ci < len(widths_cm):
                    cell.width = Cm(widths_cm[ci])
