"""HUB 合同模板 docx 占位符解析。

从 docx 文件字节流中扫描所有 {{name}} 形式的占位符，支持嵌套表格和中文占位符名。
"""
from __future__ import annotations

import io
import re

from docx import Document
from fastapi import HTTPException

# 默认占位符中文显示名映射 — 上传时自动填,admin 后台可改。
# 设计原则：用户在钉钉 / admin UI 看到的是中文标签（"客户名"）,docx 模板里写的还是英文 code。
# 漏在字典里的 code → label 默认 = code 本身,admin 自己给中文名。
_DEFAULT_LABELS: dict[str, str] = {
    # 客户类
    "customer_id": "客户 ID",
    "customer_name": "客户名",
    "customer_address": "客户地址",
    "customer_phone": "客户电话",
    "customer_tax_id": "客户税号",
    # 收货 / 联系人
    "shipping_address": "收货地址",
    "shipping_contact": "联系人",
    "shipping_phone": "联系电话",
    # 商品 / 金额
    "items": "商品明细",
    "products": "商品明细",
    "total_amount": "总金额",
    "total": "总金额",
    "subtotal": "小计",
    "tax_rate": "税率",
    "tax_amount": "税额",
    # 合同元信息
    "contract_no": "合同号",
    "contract_number": "合同号",
    "contract_date": "签订日期",
    "sign_date": "签订日期",
    "effective_date": "生效日期",
    "payment_terms": "付款方式",
    "payment_method": "付款方式",
    # 销售方（卖方）
    "seller_name": "卖方名称",
    "seller_address": "卖方地址",
    "seller_phone": "卖方电话",
    "seller_tax_id": "卖方税号",
    # 报价单
    "quote_no": "报价单号",
    "quote_date": "报价日期",
    "valid_until": "有效期至",
    # 备注
    "remarks": "备注",
    "notes": "备注",
}


def _label_for(name: str) -> str:
    """name → 中文显示名。优先字典命中,否则原样返（admin 在 UI 改成中文）。"""
    return _DEFAULT_LABELS.get(name, name)


def _enrich_placeholders(placeholders: list[dict] | None) -> list[dict]:
    """老模板 placeholders 字段没 label 字段(v1 版本上传的) → 懒填默认 label。
    新模板 upload 时已经写好 label,这里幂等通过。
    """
    out = []
    for ph in placeholders or []:
        if not isinstance(ph, dict):
            continue
        item = dict(ph)  # 不 mutate 原 dict
        if not item.get("label"):
            item["label"] = _label_for(item.get("name", ""))
        out.append(item)
    return out


def _scan_cell(cell, found: dict, pattern) -> None:
    """递归扫 cell 的段落 + 嵌套表格。

    v2 加固（review M1）：支持嵌套表格递归扫描。
    """
    for para in cell.paragraphs:
        for m in pattern.finditer(para.text):
            name = m.group(1)
            if name not in found:
                found[name] = {
                    "name": name,
                    "label": _label_for(name),
                    "type": "string",
                    "required": True,
                }
    for nested_table in cell.tables:
        for row in nested_table.rows:
            for nested_cell in row.cells:
                _scan_cell(nested_cell, found, pattern)


def _extract_placeholders(docx_bytes: bytes) -> list[dict]:
    """从 docx 文件字节流解析所有 {{name}} 形式的占位符。

    - 扫描所有段落（paragraphs）和表格单元格（tables），含嵌套表格
    - 支持中文占位符名（re.UNICODE）
    - 返回 list[{"name": str, "type": "string", "required": True}]（第一版默认 type=string）
    - 同名占位符只保留一条（去重）

    v2 加固（review M1）：pattern 加 re.UNICODE，递归嵌套表格。
    """
    try:
        doc = Document(io.BytesIO(docx_bytes))
    except Exception as exc:
        raise HTTPException(status_code=400, detail=f"docx 文件解析失败：{exc}") from exc

    found: dict[str, dict] = {}
    # re.UNICODE 支持中文占位符名，如 {{客户名}}
    pattern = re.compile(r"\{\{(\w+)\}\}", re.UNICODE)

    # 顶层段落
    for para in doc.paragraphs:
        for m in pattern.finditer(para.text):
            name = m.group(1)
            if name not in found:
                found[name] = {
                    "name": name,
                    "label": _label_for(name),  # 中文显示名（admin 可改）
                    "type": "string",
                    "required": True,
                }

    # 表格内段落（含嵌套表格，递归）
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                _scan_cell(cell, found, pattern)

    return list(found.values())
