"""HUB 后台合同模板管理（Plan 6 Task 11）。

提供 5 类 endpoint：
  - POST   /admin/contract-templates           上传 docx 模板（自动解析占位符）
  - GET    /admin/contract-templates           列表查询
  - GET    /admin/contract-templates/{id}/placeholders  获取已存占位符
  - PUT    /admin/contract-templates/{id}      更新元信息（不重传文件）
  - POST   /admin/contract-templates/{id}/disable     停用
  - POST   /admin/contract-templates/{id}/enable      启用
"""
from __future__ import annotations

import base64
import io
import logging
import re
from typing import Literal

from docx import Document
from fastapi import APIRouter, Depends, File, Form, HTTPException, Query, Request, UploadFile
from pydantic import BaseModel, Field

from hub.auth.admin_perms import require_hub_perm
from hub.models.audit import AuditLog
from hub.models.contract import ContractTemplate

logger = logging.getLogger("hub.routers.admin.contract_templates")

router = APIRouter(prefix="/hub/v1/admin/contract-templates", tags=["admin", "contract_templates"])

# 合同模板管理权限码（Plan 6 Task 11 新增，见 seed.py）
_PERM_WRITE = "usecase.contract_templates.write"
# v2 加固（review C1）：读权限独立，方便 viewer 角色只读
_PERM_READ = "usecase.contract_templates.read"

# 允许的模板类型
_VALID_TYPES = {"sales", "purchase", "framework", "quote", "other"}

# 文件大小上限：5MB
_MAX_FILE_SIZE = 5 * 1024 * 1024


# 默认占位符中文显示名映射 — 上传时自动填,admin 后台可改。
# 设计原则：用户在钉钉 / admin UI 看到的是中文标签（"客户名"）,docx 模板里写的还是英文 code。
# 漏在字典里的 code → label 默认 = code 本身,admin 自己给中文名。
_DEFAULT_LABELS: dict[str, str] = {
    # 客户类
    "customer_id": "客户 ID",
    "customer_name": "客户名",
    "customer_address": "客户地址",
    "customer_phone": "客户电话",
    "customer_tax_id": "客户税号",
    # 收货 / 联系人
    "shipping_address": "收货地址",
    "shipping_contact": "联系人",
    "shipping_phone": "联系电话",
    # 商品 / 金额
    "items": "商品明细",
    "products": "商品明细",
    "total_amount": "总金额",
    "total": "总金额",
    "subtotal": "小计",
    "tax_rate": "税率",
    "tax_amount": "税额",
    # 合同元信息
    "contract_no": "合同号",
    "contract_number": "合同号",
    "contract_date": "签订日期",
    "sign_date": "签订日期",
    "effective_date": "生效日期",
    "payment_terms": "付款方式",
    "payment_method": "付款方式",
    # 销售方（卖方）
    "seller_name": "卖方名称",
    "seller_address": "卖方地址",
    "seller_phone": "卖方电话",
    "seller_tax_id": "卖方税号",
    # 报价单
    "quote_no": "报价单号",
    "quote_date": "报价日期",
    "valid_until": "有效期至",
    # 备注
    "remarks": "备注",
    "notes": "备注",
}


def _label_for(name: str) -> str:
    """name → 中文显示名。优先字典命中,否则原样返（admin 在 UI 改成中文）。"""
    return _DEFAULT_LABELS.get(name, name)


def _enrich_placeholders(placeholders: list[dict] | None) -> list[dict]:
    """老模板 placeholders 字段没 label 字段(v1 版本上传的) → 懒填默认 label。
    新模板 upload 时已经写好 label,这里幂等通过。
    """
    out = []
    for ph in placeholders or []:
        if not isinstance(ph, dict):
            continue
        item = dict(ph)  # 不 mutate 原 dict
        if not item.get("label"):
            item["label"] = _label_for(item.get("name", ""))
        out.append(item)
    return out


class ContractTemplateRow(BaseModel):
    id: int
    name: str
    template_type: str
    description: str | None = None
    placeholders: list[dict]
    is_active: bool
    created_by_hub_user_id: int | None = None
    created_at: str


class ContractTemplateUpdate(BaseModel):
    name: str | None = Field(default=None, max_length=200)
    # v2 加固（review M2）：用 Literal 替代手工校验
    template_type: Literal["sales", "purchase", "framework", "quote", "other"] | None = None
    description: str | None = Field(default=None, max_length=1000)


def _scan_cell(cell, found: dict, pattern) -> None:
    """递归扫 cell 的段落 + 嵌套表格。

    v2 加固（review M1）：支持嵌套表格递归扫描。
    """
    for para in cell.paragraphs:
        for m in pattern.finditer(para.text):
            name = m.group(1)
            if name not in found:
                found[name] = {
                    "name": name,
                    "label": _label_for(name),
                    "type": "string",
                    "required": True,
                }
    for nested_table in cell.tables:
        for row in nested_table.rows:
            for nested_cell in row.cells:
                _scan_cell(nested_cell, found, pattern)


def _extract_placeholders(docx_bytes: bytes) -> list[dict]:
    """从 docx 文件字节流解析所有 {{name}} 形式的占位符。

    - 扫描所有段落（paragraphs）和表格单元格（tables），含嵌套表格
    - 支持中文占位符名（re.UNICODE）
    - 返回 list[{"name": str, "type": "string", "required": True}]（第一版默认 type=string）
    - 同名占位符只保留一条（去重）

    v2 加固（review M1）：pattern 加 re.UNICODE，递归嵌套表格。
    """
    try:
        doc = Document(io.BytesIO(docx_bytes))
    except Exception as exc:
        raise HTTPException(status_code=400, detail=f"docx 文件解析失败：{exc}") from exc

    found: dict[str, dict] = {}
    # re.UNICODE 支持中文占位符名，如 {{客户名}}
    pattern = re.compile(r"\{\{(\w+)\}\}", re.UNICODE)

    # 顶层段落
    for para in doc.paragraphs:
        for m in pattern.finditer(para.text):
            name = m.group(1)
            if name not in found:
                found[name] = {
                    "name": name,
                    "label": _label_for(name),  # 中文显示名（admin 可改）
                    "type": "string",
                    "required": True,
                }

    # 表格内段落（含嵌套表格，递归）
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                _scan_cell(cell, found, pattern)

    return list(found.values())


@router.post(
    "",
    dependencies=[Depends(require_hub_perm(_PERM_WRITE))],
    summary="上传合同模板",
)
async def upload_template(
    request: Request,
    name: str = Form(..., max_length=200),
    template_type: str = Form(..., max_length=50),
    # v2 加固（review M6）：description 改为 Optional，允许不传
    description: str | None = Form(None, max_length=1000),
    file: UploadFile = File(...),
):
    """上传合同模板 docx 文件。

    - 文件大小上限 5MB，格式必须为 .docx
    - 自动解析 {{xxx}} 占位符，第一版默认 type=string + required=True
    - file_storage_key 第一版用 base64 编码存储（与 Task 7 ContractRenderer 对齐）
    """
    # 1. 基础校验
    if not name.strip():
        raise HTTPException(status_code=400, detail="模板名称不能为空")
    if template_type not in _VALID_TYPES:
        raise HTTPException(status_code=400, detail="模板类型必须是 sales/purchase/framework/quote/other")
    if not file.filename or not file.filename.lower().endswith(".docx"):
        raise HTTPException(status_code=400, detail="仅支持 .docx 格式的文件")

    # v2 加固（review I1）：stream-read 64KB 分块提前 abort，防大文件 DoS
    if file.size and file.size > _MAX_FILE_SIZE:
        raise HTTPException(status_code=400, detail=f"文件超过 {_MAX_FILE_SIZE // (1024 * 1024)}MB 上限，请压缩后重试")

    buf = bytearray()
    while chunk := await file.read(64 * 1024):
        buf.extend(chunk)
        if len(buf) > _MAX_FILE_SIZE:
            raise HTTPException(status_code=400, detail=f"文件超过 {_MAX_FILE_SIZE // (1024 * 1024)}MB 上限，请压缩后重试")
    content = bytes(buf)

    if not content:
        raise HTTPException(status_code=400, detail="文件内容为空，请重新上传")

    # 2. 自动解析占位符
    placeholders = _extract_placeholders(content)

    # 3. 持久化
    actor = request.state.hub_user
    file_storage_key = base64.b64encode(content).decode("ascii")

    # v2 加固（review M6）：description 用 or None 语义更清晰
    description_clean = (description or "").strip() or None

    template = await ContractTemplate.create(
        name=name.strip(),
        template_type=template_type,
        file_storage_key=file_storage_key,
        placeholders=placeholders,
        description=description_clean,
        is_active=True,
        created_by_hub_user_id=actor.id,
    )

    # v2 加固（review C2）：合同模板为法律级文件，写操作必须留审计
    await AuditLog.create(
        who_hub_user_id=actor.id,
        action="upload_contract_template",
        target_type="contract_template",
        target_id=str(template.id),
        detail={
            "name": template.name,
            "template_type": template.template_type,
            "placeholders_count": len(placeholders),
            "file_size_bytes": len(content),
        },
    )

    logger.info("合同模板上传成功 id=%d name=%s placeholders=%d", template.id, template.name, len(placeholders))
    return {
        "id": template.id,
        "name": template.name,
        "template_type": template.template_type,
        "placeholders": placeholders,
        "is_active": True,
        "message": f"上传成功，共识别 {len(placeholders)} 个占位符",
    }


@router.get(
    "",
    # v2 加固（review C1）：GET list 也需鉴权，任何登录用户不能随意查看模板内容
    dependencies=[Depends(require_hub_perm(_PERM_READ))],
    summary="合同模板列表",
)
async def list_templates(
    template_type: str | None = None,
    is_active: bool | None = None,
    # v2 加固（review M5）：limit/offset 加 clamp，防无限查询
    limit: int = Query(100, ge=1, le=500),
    offset: int = Query(0, ge=0),
):
    """列出所有合同模板（管理员可见全部）。

    支持按模板类型和启用状态筛选。
    """
    qs = ContractTemplate.all().order_by("-created_at")
    if template_type:
        qs = qs.filter(template_type=template_type)
    if is_active is not None:
        qs = qs.filter(is_active=is_active)

    total = await qs.count()
    rows = await qs.offset(offset).limit(limit).all()

    return {
        "items": [
            ContractTemplateRow(
                id=r.id,
                name=r.name,
                template_type=r.template_type,
                description=r.description,
                placeholders=_enrich_placeholders(r.placeholders),
                is_active=r.is_active,
                created_by_hub_user_id=r.created_by_hub_user_id,
                created_at=r.created_at.isoformat(),
            ).model_dump()
            for r in rows
        ],
        "total": total,
    }


@router.get(
    "/{template_id}/placeholders",
    # v2 加固（review C1）：GET placeholders 也需鉴权
    dependencies=[Depends(require_hub_perm(_PERM_READ))],
    summary="查看模板占位符",
)
async def get_placeholders(template_id: int):
    """获取已解析的占位符列表。

    Plan 偏离说明：plan §3055 字面要求"解析 docx 找出所有 {{xxx}}"，
    实现简化为读 DB 的 placeholders JSON 列（upload 时已解析并持久化）。
    优势：避免每次查询都 base64-decode + python-docx 解析（性能 ~10x）。
    限制：当前没有"重新上传 docx 替换文件"的 endpoint，所以 DB 里的 placeholders
    不会过期；如果未来加重传 endpoint，需同步刷新此列。

    v2 加固（review I2）：docstring 说明 plan 偏离原因。
    """
    template = await ContractTemplate.filter(id=template_id).first()
    if not template:
        raise HTTPException(status_code=404, detail=f"模板 {template_id} 不存在")
    # _enrich_placeholders 处理 v1 老模板 placeholders 字段缺 label 字段的情况,
    # 用 _DEFAULT_LABELS 字典懒填中文显示名（admin 之后可改）。
    return {"placeholders": _enrich_placeholders(template.placeholders)}


@router.put(
    "/{template_id}/placeholders",
    dependencies=[Depends(require_hub_perm(_PERM_WRITE))],
    summary="保存占位符 label / type / required（admin 编辑后调）",
)
async def update_placeholders(
    request: Request,
    template_id: int,
    body: dict = None,  # noqa: B008  — pydantic 严格 schema 见下面校验
):
    """保存 admin 编辑后的占位符列表。

    Body schema: `{"placeholders": [{"name": str, "label": str, "type": str, "required": bool}, ...]}`

    安全校验:
    - 不允许新增/删除占位符（占位符列表由 docx 解析得出,只能改 metadata）
    - 只允许改 label / type / required;name 不可改（动 name 等于改 docx 里的 {{xxx}},不一致）
    - 必须给齐 placeholders 列表（保留原 set 即可,不允许漏报）
    """
    template = await ContractTemplate.filter(id=template_id).first()
    if not template:
        raise HTTPException(status_code=404, detail=f"模板 {template_id} 不存在")

    incoming = (body or {}).get("placeholders")
    if not isinstance(incoming, list):
        raise HTTPException(status_code=400, detail="placeholders 必须是列表")

    existing = _enrich_placeholders(template.placeholders)
    existing_names = {p["name"] for p in existing}
    incoming_names = {p.get("name") for p in incoming if isinstance(p, dict)}
    if existing_names != incoming_names:
        raise HTTPException(
            status_code=400,
            detail=(
                f"占位符 name 集合必须跟模板原始解析结果一致。"
                f"缺失: {sorted(existing_names - incoming_names)};"
                f"多余: {sorted(incoming_names - existing_names)}"
            ),
        )

    # 校验后规范化（防 admin 漏 / 乱传字段）
    valid_types = {"string", "number", "date", "money", "phone"}
    cleaned = []
    for p in incoming:
        name = p.get("name") or ""
        label = (p.get("label") or "").strip() or _label_for(name)
        ptype = p.get("type") or "string"
        if ptype not in valid_types:
            ptype = "string"
        required = bool(p.get("required", True))
        cleaned.append({
            "name": name,
            "label": label,
            "type": ptype,
            "required": required,
        })

    template.placeholders = cleaned
    await template.save()

    actor = request.state.hub_user
    await AuditLog.create(
        who_hub_user_id=actor.id,
        action="update_template_placeholders",
        target_type="contract_template",
        target_id=str(template_id),
        detail={"count": len(cleaned)},
    )
    return {"success": True, "placeholders": cleaned}


@router.get(
    "/{template_id}/file",
    dependencies=[Depends(require_hub_perm(_PERM_READ))],
    summary="下载模板 docx 原文件",
)
async def get_template_file(template_id: int):
    """返回 docx 原文件字节流,前端用 mammoth.js 转 HTML 做可视化预览。

    file_storage_key 当前是 base64 编码的 docx,decode 后用 StreamingResponse
    返。前端 fetch 这个 endpoint 拿 ArrayBuffer → mammoth.convertToHtml() → 渲染。
    """
    from fastapi.responses import Response

    template = await ContractTemplate.filter(id=template_id).first()
    if not template:
        raise HTTPException(status_code=404, detail=f"模板 {template_id} 不存在")
    if not template.file_storage_key:
        raise HTTPException(status_code=404, detail="模板文件内容为空")
    try:
        content = base64.b64decode(template.file_storage_key)
    except Exception as exc:
        raise HTTPException(status_code=500, detail=f"模板文件 base64 解码失败：{exc}") from exc
    return Response(
        content=content,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={
            # 用 Content-Disposition 给文件名,但 inline 而非 attachment（前端 fetch 不下载）
            "Content-Disposition": f'inline; filename="template-{template_id}.docx"',
        },
    )


@router.put(
    "/{template_id}",
    dependencies=[Depends(require_hub_perm(_PERM_WRITE))],
    summary="更新模板元信息",
)
async def update_template(request: Request, template_id: int, body: ContractTemplateUpdate):
    """更新模板名称、类型或描述（不支持重新上传文件，不更新占位符）。"""
    template = await ContractTemplate.filter(id=template_id).first()
    if not template:
        raise HTTPException(status_code=404, detail=f"模板 {template_id} 不存在")

    if body.name is not None:
        if not body.name.strip():
            raise HTTPException(status_code=400, detail="模板名称不能为空")
        template.name = body.name.strip()
    if body.template_type is not None:
        # Literal 已在 Pydantic 层校验，此处无需手工校验
        template.template_type = body.template_type
    if body.description is not None:
        template.description = body.description.strip() or None

    await template.save()

    # v2 加固（review C2）：更新元信息写审计
    await AuditLog.create(
        who_hub_user_id=request.state.hub_user.id,
        action="update_contract_template",
        target_type="contract_template",
        target_id=str(template_id),
        detail={"changes": body.model_dump(exclude_none=True)},
    )

    return {"id": template.id, "message": "模板信息已更新"}


@router.post(
    "/{template_id}/disable",
    dependencies=[Depends(require_hub_perm(_PERM_WRITE))],
    summary="停用模板",
)
async def disable_template(request: Request, template_id: int):
    """将模板标记为禁用（is_active=False）。已禁用时幂等返回。"""
    template = await ContractTemplate.filter(id=template_id).first()
    if not template:
        raise HTTPException(status_code=404, detail=f"模板 {template_id} 不存在")
    if not template.is_active:
        return {"id": template.id, "message": "该模板已是禁用状态"}
    template.is_active = False
    await template.save()

    # v2 加固（review C2）：停用写审计
    await AuditLog.create(
        who_hub_user_id=request.state.hub_user.id,
        action="disable_contract_template",
        target_type="contract_template",
        target_id=str(template_id),
        detail={"name": template.name},
    )

    return {"id": template.id, "message": "已禁用"}


@router.post(
    "/{template_id}/enable",
    dependencies=[Depends(require_hub_perm(_PERM_WRITE))],
    summary="启用模板",
)
async def enable_template(request: Request, template_id: int):
    """将模板标记为启用（is_active=True）。已启用时幂等返回。"""
    template = await ContractTemplate.filter(id=template_id).first()
    if not template:
        raise HTTPException(status_code=404, detail=f"模板 {template_id} 不存在")
    if template.is_active:
        return {"id": template.id, "message": "该模板已是启用状态"}
    template.is_active = True
    await template.save()

    # v2 加固（review C2）：启用写审计
    await AuditLog.create(
        who_hub_user_id=request.state.hub_user.id,
        action="enable_contract_template",
        target_type="contract_template",
        target_id=str(template_id),
        detail={"name": template.name},
    )

    return {"id": template.id, "message": "已启用"}
